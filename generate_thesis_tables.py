"""
Generates thesis_tables.docx with all formal-paper tables for the diploma thesis.

Real numbers come from:
  - results/benchmark_runs.csv  (aggregated per scope/mode)
  - batch_requests/*.ps1        (dataset sizes)
  - backend/routes/*.py         (API endpoints, algorithm constants)
"""

import csv
import statistics as st
from collections import defaultdict
from pathlib import Path

from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


ROOT = Path("/Users/mateogrifsha/Documents/mateo/uni_auto_scheduler")
OUT  = ROOT / "thesis_tables.docx"


# ---------- doc styling ----------

def set_cell_border(cell, top=False, bottom=False, color="000000", size="6"):
    """Set only the requested horizontal borders; suppress all others (three-line table style)."""
    tc_pr = cell._tc.get_or_add_tcPr()
    # Remove any pre-existing border element so our setting is authoritative
    existing = tc_pr.find(qn("w:tcBorders"))
    if existing is not None:
        tc_pr.remove(existing)
    borders = OxmlElement("w:tcBorders")
    for edge in ("top", "left", "bottom", "right"):
        b = OxmlElement(f"w:{edge}")
        if (edge == "top" and top) or (edge == "bottom" and bottom):
            b.set(qn("w:val"), "single")
            b.set(qn("w:sz"), size)
            b.set(qn("w:color"), color)
        else:
            b.set(qn("w:val"), "nil")
        borders.append(b)
    tc_pr.append(borders)


def shade_cell(cell, hex_color="D9D9D9"):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    tc_pr.append(shd)


def style_run(run, bold=False, size=12, font="Times New Roman"):
    run.font.name = font
    run.font.size = Pt(size)
    run.bold = bold
    # Force the font on East Asian / complex script slots too so Word does not substitute
    rPr = run._element.get_or_add_rPr()
    rFonts = rPr.find(qn("w:rFonts"))
    if rFonts is None:
        rFonts = OxmlElement("w:rFonts")
        rPr.append(rFonts)
    for attr in ("w:ascii", "w:hAnsi", "w:cs", "w:eastAsia"):
        rFonts.set(qn(attr), font)


def add_caption(doc, text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(text)
    style_run(r, bold=True, size=12)
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(2)


def add_section(doc, text):
    p = doc.add_paragraph()
    r = p.add_run(text)
    style_run(r, bold=True, size=14)
    p.paragraph_format.space_before = Pt(14)
    p.paragraph_format.space_after = Pt(4)


def add_table(doc, headers, rows, numeric_cols=None, col_widths_cm=None):
    """Build a three-line academic table: top rule above header, rule under header, bottom rule under last row."""
    numeric_cols = numeric_cols or set()
    n_rows = 1 + len(rows)
    table = doc.add_table(rows=n_rows, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"  # we override borders cell-by-cell below

    # Header row
    for j, h in enumerate(headers):
        cell = table.rows[0].cells[j]
        cell.text = ""
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(str(h))
        style_run(r, bold=True, size=12)
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        set_cell_border(cell, top=True, bottom=True)
        if col_widths_cm and j < len(col_widths_cm):
            cell.width = Cm(col_widths_cm[j])

    # Body rows
    for i, row in enumerate(rows, start=1):
        is_last = (i == n_rows - 1)
        for j, val in enumerate(row):
            cell = table.rows[i].cells[j]
            cell.text = ""
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER if j in numeric_cols else WD_ALIGN_PARAGRAPH.LEFT
            r = p.add_run(str(val))
            style_run(r, bold=False, size=12)
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            set_cell_border(cell, top=False, bottom=is_last)
            if col_widths_cm and j < len(col_widths_cm):
                cell.width = Cm(col_widths_cm[j])

    doc.add_paragraph()  # spacer
    return table


# ---------- data loaders ----------

def load_benchmark():
    rows = []
    with open(ROOT / "results" / "benchmark_runs.csv") as f:
        for r in csv.DictReader(f):
            r["total_ms"]      = float(r["total_ms"])
            r["algorithm_ms"]  = float(r["algorithm_ms"])
            r["auth_ms"]       = float(r["auth_ms"])
            r["db_read_ms"]    = float(r["db_read_ms"])
            r["db_write_ms"]   = float(r["db_write_ms"])
            r["db_rollback_ms"]= float(r["db_rollback_ms"])
            r["schedules_simulated"] = int(r["schedules_simulated"])
            r["skipped_count"] = int(r["skipped_count"])
            r["total_records"] = int(r["total_records"])
            rows.append(r)
    return rows


def agg(values):
    if not values:
        return ("-", "-", "-", "-", "-")
    return (
        f"{st.mean(values):.1f}",
        f"{st.median(values):.1f}",
        f"{(st.stdev(values) if len(values) > 1 else 0):.1f}",
        f"{min(values):.1f}",
        f"{max(values):.1f}",
    )


# ---------- main ----------

def main():
    bench = load_benchmark()
    doc = Document()

    # Default body font
    style = doc.styles["Normal"]
    style.font.name = "Times New Roman"
    style.font.size = Pt(12)

    # Title page heading
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = title.add_run("Appendix: Tables")
    style_run(r, bold=True, size=14)

    # ===== Section 1: Problem Description =====
    add_section(doc, "1. Problem Description")

    # Table 1: dataset entity counts
    add_caption(doc, "Table 1. Test dataset entity counts.")
    add_table(
        doc,
        headers=["Entity", "Count", "Source / Notes"],
        rows=[
            ["Faculties",            "2",   "Engineering, Economics"],
            ["Degrees",              "11",  "6 Engineering, 5 Economics"],
            ["Student groups",       "66",  "11 degrees x 3 years x 2 semesters"],
            ["Students",             "528", "8 students per group"],
            ["Courses (catalogue)",  "97",  "Across all degrees and years"],
            ["Professors",           "45",  "Many-to-many with courses"],
            ["Rooms",                "22",  "Capacities 10 to 150"],
            ["Time slots",           "25",  "5 days x 5 two-hour slots"],
            ["Room-slot combinations", "550", "Total scheduling capacity"],
        ],
        numeric_cols={1},
        col_widths_cm=[5.0, 2.0, 8.0],
    )

    # Table 2: curriculum distribution
    add_caption(doc, "Table 2. Curriculum distribution by year and semester (full-school run).")
    add_table(
        doc,
        headers=["Year", "Semester", "Curriculum entries", "Groups affected", "Schedules to create"],
        rows=[
            ["1", "1", "11", "11", "55"],
            ["1", "2", "11", "11", "55"],
            ["2", "1", "11", "11", "55"],
            ["2", "2", "11", "11", "55"],
            ["3", "1", "11", "11", "55"],
            ["3", "2", "11", "11", "55"],
            ["Total", "", "66", "66", "330"],
        ],
        numeric_cols={0, 1, 2, 3, 4},
        col_widths_cm=[2.0, 2.5, 3.5, 3.5, 4.0],
    )

    # Table 3: constraints
    add_caption(doc, "Table 3. Scheduling constraints by type and affected entity.")
    add_table(
        doc,
        headers=["Constraint", "Type", "Entity affected", "Enforcement"],
        rows=[
            ["No two sessions share a (professor, slot)",   "Hard", "Professor",            "In-memory set check"],
            ["No two sessions share a (group, slot)",       "Hard", "Student group",        "In-memory set check"],
            ["No two sessions share a (room, slot)",        "Hard", "Room",                 "In-memory set check (combinable for same course)"],
            ["Room capacity >= group size",                 "Hard", "Room, group",          "Filter inside find_slot_and_room"],
            ["Each course has one Lecture and one Seminar", "Hard", "Course schedule",      "Two CourseSession rows per CourseSchedule"],
            ["Lecture and Seminar on different days",       "Soft", "Group",                "exclude_day with same_day_penalty"],
            ["Professor availability window covers slot",   "Soft", "Professor",            "Open-world; blocks slots when explicit window exists"],
            ["Max courses per group (default 10)",          "Soft", "Group",                "max_courses_per_group parameter"],
            ["Day fill prefers 4-hour student days",        "Soft", "Group",                "day_session_score heuristic"],
            ["Heavier courses spread across days",          "Soft", "Group",                "c_difficulty_weight, currently unused"],
        ],
        numeric_cols=set(),
        col_widths_cm=[6.0, 1.5, 3.5, 4.5],
    )

    # ===== Section 2: Algorithm =====
    add_section(doc, "2. Algorithm")

    # Table 4: day_session_score
    add_caption(doc, "Table 4. day_session_score heuristic values.")
    add_table(
        doc,
        headers=["Existing sessions on day", "Score", "Effect"],
        rows=[
            ["0",  "1",  "Opens a new day (acceptable)"],
            ["1",  "0",  "Best: completes a 4-hour student day"],
            ["2",  "2",  "Discouraged: produces a 6-hour day"],
            ["3+", "99", "Hard avoid: would exceed 6 hours"],
        ],
        numeric_cols={0, 1},
        col_widths_cm=[5.0, 2.0, 8.0],
    )

    # Table 5: baseline vs full
    add_caption(doc, "Table 5. Comparison of baseline and full scheduling modes.")
    add_table(
        doc,
        headers=["Feature", "Baseline", "Full"],
        rows=[
            ["Professor selection",            "First in deterministic order",       "Load-balanced (min sessions assigned)"],
            ["Slot selection",                 "First valid slot",                   "Scored candidates, lowest score wins"],
            ["day_session_score heuristic",    "Disabled",                            "Enabled"],
            ["same_day_penalty (lec vs sem)",  "Soft preference only",                "5.0 + 3.0 x difficulty_weight"],
            ["Combined-lecture room sharing",  "Implicit, no bonus",                  "combine_bonus = -2.0"],
            ["Difficulty-weighted spreading",  "Disabled",                            "Weight penalty / 10.0"],
            ["Fallback when no candidate",     "Same-day allowed",                    "Same-day or full-day allowed"],
            ["Skipped on full-school run",     "81 of 330",                           "0 of 330"],
        ],
        numeric_cols=set(),
        col_widths_cm=[5.0, 5.0, 5.0],
    )

    # Table 6: parameters
    add_caption(doc, "Table 6. Tunable parameters of the scheduling algorithm.")
    add_table(
        doc,
        headers=["Parameter", "Default", "Where used"],
        rows=[
            ["max_courses_per_group",       "10",                    "Per-group cap on offerings"],
            ["mode",                        "\"full\"",               "Selects baseline or full path"],
            ["same_day_penalty (base)",     "5.0",                   "find_slot_and_room"],
            ["same_day_penalty (per weight)", "+3.0 x weight",        "find_slot_and_room"],
            ["weight_penalty divisor",      "10.0",                  "Day load weighting"],
            ["combine_bonus",               "-2.0",                  "Same-course room reuse"],
            ["Hard day cap",                "3 sessions per day",    "day_session_score returns 99"],
            ["Lecture duration",            "2 hours",               "One per CourseSchedule"],
            ["Seminar duration",            "2 hours",               "One per CourseSchedule"],
        ],
        numeric_cols={1},
        col_widths_cm=[5.5, 4.0, 5.5],
    )

    # ===== Section 3: Results / Benchmarks =====
    add_section(doc, "3. Benchmark Results")

    # Group benchmark data by scope+mode and by dataset size
    by_key = defaultdict(list)
    for r in bench:
        by_key[(r["scope"], r["mode"], r["total_records"])].append(r)

    # Pick the larger dataset (3487) as the canonical results table; show all 4 modes
    target_records = 3487
    scope_order = [("single", "Single semester (Year 1, Sem 1)"), ("full_school", "Full school (all years)")]
    mode_order  = [
        ("baseline", "Baseline"),
        ("full", "Full"),
        ("baseline_realistic", "Baseline + auth + rollback"),
        ("full_realistic", "Full + auth + rollback"),
    ]

    # Table 7: total runtime
    add_caption(doc, f"Table 7. End-to-end runtime per scheduling run, in milliseconds (n=11 per cell, dataset size {target_records} DB records).")
    rows = []
    for scope_key, scope_label in scope_order:
        for mode_key, mode_label in mode_order:
            samples = [r["total_ms"] for r in by_key.get((scope_key, mode_key, target_records), [])]
            mean, median, stdev, lo, hi = agg(samples)
            rows.append([scope_label, mode_label, mean, median, stdev, lo, hi])
    add_table(
        doc,
        headers=["Scope", "Mode", "Mean", "Median", "Std. dev.", "Min", "Max"],
        rows=rows,
        numeric_cols={2, 3, 4, 5, 6},
        col_widths_cm=[4.5, 4.0, 1.6, 1.6, 1.8, 1.4, 1.4],
    )

    # Table 8: timing breakdown
    add_caption(doc, f"Table 8. Mean time spent per phase (ms), dataset size {target_records}.")
    rows = []
    for scope_key, scope_label in scope_order:
        for mode_key, mode_label in mode_order:
            runs = by_key.get((scope_key, mode_key, target_records), [])
            if not runs:
                continue
            algo = st.mean(r["algorithm_ms"]  for r in runs)
            auth = st.mean(r["auth_ms"]       for r in runs)
            rd   = st.mean(r["db_read_ms"]    for r in runs)
            wr   = st.mean(r["db_write_ms"]   for r in runs)
            rb   = st.mean(r["db_rollback_ms"] for r in runs)
            total= st.mean(r["total_ms"]      for r in runs)
            rows.append([
                scope_label, mode_label,
                f"{algo:.1f}", f"{auth:.2f}", f"{rd:.1f}",
                f"{wr:.1f}", f"{rb:.2f}", f"{total:.1f}",
            ])
    add_table(
        doc,
        headers=["Scope", "Mode", "Algorithm", "Auth", "DB read", "DB write", "Rollback", "Total"],
        rows=rows,
        numeric_cols={2, 3, 4, 5, 6, 7},
        col_widths_cm=[3.5, 3.5, 1.6, 1.4, 1.6, 1.6, 1.6, 1.6],
    )

    # Table 9: schedules created vs skipped
    add_caption(doc, "Table 9. Schedules produced per mode (full-school scope).")
    rows = []
    for mode_key, mode_label in mode_order:
        runs = by_key.get(("full_school", mode_key, target_records), [])
        if not runs:
            continue
        sims     = runs[0]["schedules_simulated"]
        skipped  = runs[0]["skipped_count"]
        created  = sims - skipped
        rate     = 100.0 * created / sims if sims else 0
        rows.append([mode_label, sims, created, skipped, f"{rate:.1f}%"])
    add_table(
        doc,
        headers=["Mode", "Attempted", "Created", "Skipped", "Success rate"],
        rows=rows,
        numeric_cols={1, 2, 3, 4},
        col_widths_cm=[5.5, 2.5, 2.5, 2.5, 2.5],
    )

    # Table 10: dataset-size scaling (compare 2797 vs 3487)
    add_caption(doc, "Table 10. Scaling with DB record count (full-school, full mode, mean total_ms).")
    rows = []
    sizes = sorted({r["total_records"] for r in bench if r["scope"] == "full_school"})
    for sz in sizes:
        runs = by_key.get(("full_school", "full", sz), [])
        if not runs:
            continue
        mean = st.mean(r["total_ms"] for r in runs)
        algo = st.mean(r["algorithm_ms"] for r in runs)
        rows.append([sz, f"{mean:.1f}", f"{algo:.1f}", len(runs)])
    add_table(
        doc,
        headers=["DB records", "Mean total (ms)", "Mean algorithm (ms)", "Samples"],
        rows=rows,
        numeric_cols={0, 1, 2, 3},
        col_widths_cm=[3.5, 4.0, 4.5, 2.5],
    )

    # Table 11: validation conflicts
    add_caption(doc, "Table 11. Validation outcomes after generation (single semester, Year 1 Sem 1).")
    add_table(
        doc,
        headers=["Check", "Baseline", "Full"],
        rows=[
            ["Schedules created",             "55", "55"],
            ["Schedules with both sessions",  "55", "55"],
            ["Room-slot conflicts",            "0",  "0"],
            ["Group-slot conflicts",           "0",  "0"],
            ["Professor-slot conflicts",       "0",  "0"],
            ["is_valid",                       "True", "True"],
        ],
        numeric_cols={1, 2},
        col_widths_cm=[6.0, 3.5, 3.5],
    )

    # ===== Section 4: Related Work =====
    add_section(doc, "4. Related Work")

    add_caption(doc, "Table 12. Feature comparison with related university scheduling approaches.")
    add_table(
        doc,
        headers=["Feature", "This work", "CSP solvers", "ILP solvers", "Genetic algorithms"],
        rows=[
            ["Approach",                "Greedy + heuristic",     "Constraint propagation", "Exact optimisation", "Population-based"],
            ["Optimality guarantee",    "No",                     "Yes (if SAT)",            "Yes",                 "No"],
            ["Runtime on full school",  "~0.5 s",                  "Seconds to minutes",      "Minutes",             "Minutes"],
            ["Soft-constraint scoring", "Weighted heuristic",      "Difficult",               "Native (objective)",  "Native (fitness)"],
            ["Web UI and REST API",     "Yes",                     "Typically no",            "Typically no",        "Typically no",
             ],
            ["Persistent multi-user",   "Yes (PostgreSQL)",        "No",                      "No",                  "No"],
            ["Re-runnable / incremental","Yes",                    "Partial",                 "No",                  "No"],
        ],
        numeric_cols=set(),
        col_widths_cm=[3.8, 3.2, 3.0, 2.6, 3.0],
    )

    # ===== Section 5: Implementation =====
    add_section(doc, "5. Implementation Reference")

    # Table 13: data model
    add_caption(doc, "Table 13. Core data model (SQLAlchemy tables in backend/models.py).")
    add_table(
        doc,
        headers=["Table", "Purpose", "Key relationships"],
        rows=[
            ["Faculty",                "Top-level academic unit",            "1:N to Degree"],
            ["Degree",                 "Programme within a faculty",         "N:1 Faculty, 1:N StudentGroup"],
            ["StudentGroup",           "Cohort sharing a schedule",          "N:1 Degree, 1:N Student"],
            ["Student",                "Individual enrolled in a group",     "N:1 StudentGroup"],
            ["Prof (User)",            "Professor with assigned courses",    "M:N Course (professor_course)"],
            ["Course",                 "Catalogue entry",                    "M:N Prof, 1:N CourseSchedule"],
            ["CourseCurriculum",       "Which course belongs to which (degree, year, semester)", "N:1 Course, N:1 Degree"],
            ["CourseSchedule",         "One scheduled offering (course + group)", "N:1 Course, N:1 StudentGroup"],
            ["CourseSession",          "Lecture or Seminar instance",        "N:1 CourseSchedule, N:1 TimeSlot, N:1 Room"],
            ["TimeSlots",              "(day, start_time, end_time) tuple",  "Referenced by CourseSession"],
            ["Rooms",                  "Physical room with capacity",        "Referenced by CourseSession"],
            ["ProfessorAvailability",  "Per-professor working window",       "Read by auto_schedule, stored only"],
            ["SessionTypeModel",       "Enum table (Lecture, Seminar)",      "Referenced by CourseSession"],
        ],
        numeric_cols=set(),
        col_widths_cm=[4.0, 6.0, 5.5],
    )

    # Table 14: API endpoints (selection)
    add_caption(doc, "Table 14. Selected REST endpoints exposed by the backend (129 total across all routers).")
    add_table(
        doc,
        headers=["Method", "Path", "Purpose"],
        rows=[
            ["POST",   "/auto-schedule/generate",     "Run scheduler for (year, semester_number)"],
            ["GET",    "/auto-schedule/validate",     "Check conflicts and completeness"],
            ["DELETE", "/auto-schedule/clear",        "Remove schedules for a semester"],
            ["GET",    "/course-schedules/",          "List existing schedules"],
            ["PUT",    "/course-schedules/{id}",      "Edit a single schedule"],
            ["GET",    "/courses/",                   "List courses"],
            ["GET",    "/professors/",                "List professors with assigned courses"],
            ["GET",    "/rooms/",                     "List rooms with capacities"],
            ["GET",    "/time-slots/",                "List configured slots"],
            ["GET",    "/student-groups/",            "List groups"],
            ["POST",   "/auth/login",                 "Issue JWT for a user"],
            ["POST",   "/professor-availability/",    "Record professor working window"],
            ["GET",    "/dashboard/stats",            "Aggregated entity counts"],
            ["POST",   "/ai-suggestions/",            "Request AI-assisted suggestion"],
            ["POST",   "/complaints/",                "File a student complaint"],
            ["POST",   "/resolutions/",               "Resolve a flagged conflict"],
            ["POST",   "/benchmark/run",              "Run benchmark suite from API"],
        ],
        numeric_cols={0},
        col_widths_cm=[2.0, 5.5, 8.0],
    )

    doc.save(OUT)
    print(f"Wrote {OUT}")


if __name__ == "__main__":
    main()
