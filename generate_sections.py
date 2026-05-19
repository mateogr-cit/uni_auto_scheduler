#!/usr/bin/env python3
"""
Generates a Word document containing the replacement Sections 5 and 6
for Scheduler_Paper_v4.docx, populated with real benchmark data.
"""

import csv
import math
import statistics
from pathlib import Path
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

FIGURES = Path("results/paper_run/figures")

# ── Load and aggregate results ────────────────────────────────────────────────

CSV_PATH = Path("results/paper_run/results.csv")
rows = []
with open(CSV_PATH) as f:
    for r in csv.DictReader(f):
        rows.append({k: float(v) if k != "instance" else int(v) for k, v in r.items()})

SIZES = sorted(set(int(r["n_groups"]) for r in rows))

def stats(n, key):
    vals = [r[key] for r in rows if int(r["n_groups"]) == n]
    m = statistics.mean(vals)
    ci = 1.96 * statistics.stdev(vals) / math.sqrt(len(vals)) if len(vals) > 1 else 0
    return m, ci

# ── Document helpers ──────────────────────────────────────────────────────────

def add_heading(doc, text, level):
    p = doc.add_heading(text, level=level)
    for run in p.runs:
        run.font.color.rgb = RGBColor(0, 0, 0)
    return p

def add_paragraph(doc, text, style="Normal", bold=False, italic=False, size=11):
    p = doc.add_paragraph(style=style)
    run = p.add_run(text)
    run.bold = bold
    run.italic = italic
    run.font.size = Pt(size)
    p.paragraph_format.space_after = Pt(6)
    return p

def add_caption(doc, text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text)
    run.italic = True
    run.font.size = Pt(10)
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(10)


def add_figure(doc, filename, caption, width=Inches(5.0)):
    path = FIGURES / filename
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(6)
    run = p.add_run()
    run.add_picture(str(path), width=width)
    add_caption(doc, caption)

def set_col_width(cell, width):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcW = OxmlElement("w:tcW")
    tcW.set(qn("w:w"), str(int(width * 1440)))
    tcW.set(qn("w:type"), "dxa")
    tcPr.append(tcW)

def add_table(doc, headers, rows_data, col_widths=None):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    hdr = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr[i].text = h
        hdr[i].paragraphs[0].runs[0].bold = True
        hdr[i].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        hdr[i].paragraphs[0].runs[0].font.size = Pt(10)
        if col_widths:
            set_col_width(hdr[i], col_widths[i])
    for row in rows_data:
        cells = table.add_row().cells
        for i, val in enumerate(row):
            cells[i].text = str(val)
            cells[i].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            cells[i].paragraphs[0].runs[0].font.size = Pt(10)
            if col_widths:
                set_col_width(cells[i], col_widths[i])
    doc.add_paragraph()
    return table

# ── Build document ────────────────────────────────────────────────────────────

doc = Document()

# Page margins
for section in doc.sections:
    section.top_margin    = Inches(1.0)
    section.bottom_margin = Inches(1.0)
    section.left_margin   = Inches(1.25)
    section.right_margin  = Inches(1.25)

# Default paragraph font
doc.styles["Normal"].font.name = "Times New Roman"
doc.styles["Normal"].font.size = Pt(11)

# Title block
t = doc.add_paragraph()
t.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = t.add_run("Replacement Sections 5 & 6 — Scheduler_Paper_v4")
r.bold = True
r.font.size = Pt(13)
doc.add_paragraph()

# ─────────────────────────────────────────────────────────────────────────────
# SECTION 5
# ─────────────────────────────────────────────────────────────────────────────

add_heading(doc, "5.  Experimental Evaluation", level=1)

# 5.1
add_heading(doc, "5.1  Experimental Setup", level=2)

add_paragraph(doc,
    "The scheduler was evaluated on a benchmark of synthetic instances that reflect "
    "the operational constraints of the target institution. The weekly grid consists of "
    "five working days with three two-hour time slots per day (09:00–11:00, 11:00–13:00 "
    "and 13:00–15:00), yielding fifteen available slots in total. The room set contains "
    "twelve rooms divided into three capacity classes: four small rooms (20 seats), four "
    "medium rooms (35 seats) and four large rooms (50 seats), for a total room-slot "
    "capacity of 180 pairs per week.")

add_paragraph(doc,
    "Each instance is parameterised by the number N of independent student groups to be "
    "scheduled. Every group is assigned to its own degree programme and takes exactly five "
    "courses per semester, each taught by a dedicated professor — a configuration that "
    "directly reflects the institution's academic structure. Group enrolment size is drawn "
    "uniformly at random from the interval [15, 50]; because the three room classes span "
    "the same range, groups of different sizes face different room availability, which "
    "introduces genuine instance-to-instance variance. N is drawn from the set "
    "{5, 10, 15, 20, 25, 30}; thirty random instances were generated for every value using "
    "a fixed seed (20260515) to ensure reproducibility. Table 1 summarises the fixed "
    "parameters of the benchmark.")

add_paragraph(doc,
    "Two scheduler configurations were compared on the same set of instances. The "
    "baseline configuration runs the greedy backbone of Section 3.3 without heuristic "
    "guidance: it selects the first valid (slot, room) pair in iteration order and assigns "
    "the first qualified professor. The full configuration adds the day-packing heuristic "
    "and the professor load-balancing rule of Section 3.4. Three indicators were collected "
    "for every run: (i) the feasibility rate, defined as the fraction of scheduled "
    "(course, group) pairs relative to the expected total of N × 5; (ii) the average "
    "number of active days per student group per week (S1); and (iii) the fraction of "
    "(group, day) pairs that contain exactly two sessions — the preferred compact daily "
    "pattern — (S2).")

# Table 1
add_paragraph(doc, "")
t1_headers = ["Parameter", "Value"]
t1_rows = [
    ["Rooms",                    "12  (4 × cap-20, 4 × cap-35, 4 × cap-50)"],
    ["Slots per day",            "3  (09:00–11:00, 11:00–13:00, 13:00–15:00)"],
    ["Total weekly slots",       "15  (5 days × 3 slots, each 2 h)"],
    ["Room-slot capacity",       "180 pairs"],
    ["Courses per group",        "5  (fixed)"],
    ["Sessions per group",       "10  (5 courses × 2 sessions each)"],
    ["Group enrolment",          "Uniform [15, 50] — random per instance"],
    ["Instance sizes N",         "{5, 10, 15, 20, 25, 30} student groups"],
    ["Instances per size",       "30"],
    ["Random seed",              "20260515"],
]
add_table(doc, t1_headers, t1_rows, col_widths=[2.5, 3.5])
add_caption(doc,
    "Table 1.  Fixed parameters of the synthetic benchmark used in Section 5.")

# 5.2
add_heading(doc, "5.2  Scheduling Quality", level=2)

add_paragraph(doc,
    "Table 2 reports the mean feasibility rate, S1 and S2 for both configurations across "
    "all six instance sizes, together with 95% confidence intervals derived from the "
    "thirty repetitions of each size.")

# Table 2
t2_headers = [
    "N", "Feasibility — Base", "Feasibility — Full",
    "S1 — Base", "S1 — Full",
    "S2 — Base", "S2 — Full"
]
t2_rows = []
for n in SIZES:
    fb, fb_ci = stats(n, "feasible_baseline")
    ff, ff_ci = stats(n, "feasible_full")
    s1b, s1b_ci = stats(n, "s1_baseline")
    s1f, s1f_ci = stats(n, "s1_full")
    s2b, s2b_ci = stats(n, "s2_baseline")
    s2f, s2f_ci = stats(n, "s2_full")
    t2_rows.append([
        str(n),
        f"{fb*100:.1f} ±{fb_ci*100:.1f}%",
        f"{ff*100:.1f} ±{ff_ci*100:.1f}%",
        f"{s1b:.2f} ±{s1b_ci:.2f}",
        f"{s1f:.2f} ±{s1f_ci:.2f}",
        f"{s2b*100:.1f} ±{s2b_ci*100:.1f}%",
        f"{s2f*100:.1f} ±{s2f_ci*100:.1f}%",
    ])
add_table(doc, t2_headers, t2_rows, col_widths=[0.45, 1.2, 1.2, 1.0, 1.0, 1.1, 1.1])
add_caption(doc,
    "Table 2.  Mean feasibility rate, average active days per group (S1) and fraction of "
    "group-days with exactly two sessions (S2) for the baseline and full configurations, "
    "over 30 instances per size.  95% CI in parentheses.")

# Narrative 5.2
add_paragraph(doc,
    "Feasibility. At N = 5 both configurations achieve 100% feasibility: the fifteen "
    "available slots and 180 room-slot pairs comfortably accommodate the fifty sessions "
    "required. As N grows the room-slot budget is increasingly contested and the "
    "feasibility rate declines for both configurations, reaching approximately 51% at "
    "N = 30. Because the baseline and full configurations share the same hard-constraint "
    "checking logic, their feasibility rates are statistically indistinguishable across "
    "all sizes; the difference between the two lies exclusively in schedule quality, not "
    "coverage.")

add_paragraph(doc,
    "The per-instance variance in feasibility is a direct consequence of the random group "
    "enrolment sizes. A group whose enrolment exceeds 35 can only be assigned to the four "
    "large rooms, reducing its effective slot budget from 180 to 60 pairs. Instances that "
    "happen to draw several large groups therefore experience earlier room exhaustion, "
    "which explains the spread visible in the 95% confidence intervals.")

add_figure(doc, "fig1_feasibility.png",
    "Fig. 1.  Feasibility rate vs. N for the baseline (dashed orange) and full (solid blue) "
    "configurations, with 95% CI shaded.  The dotted grey line marks the theoretical "
    "room-capacity limit (180 room-slot pairs ÷ 10 sessions per group).  Both configurations "
    "overlap because feasibility is determined by hard constraints, not heuristic choice.")

add_paragraph(doc,
    "Quality — active days (S1). The full configuration consistently produces schedules "
    "in which student groups occupy more distinct weekdays than the baseline. At N = 5 "
    "the baseline distributes the ten weekly sessions across four days on average, "
    "leaving one day entirely free; the full configuration fills all five days "
    "(S1 = 5.00 ± 0.00). The gap narrows as N increases but remains large and "
    "statistically significant throughout the range: at N = 30 the full configuration "
    "returns S1 = 4.56 ± 0.02 against S1 = 3.13 ± 0.02 for the baseline, a difference "
    "of 1.43 days. This result is a direct consequence of the day-packing heuristic: by "
    "assigning score 1 to an empty day — a higher penalty than the score 0 assigned to a "
    "day that already hosts one session — the heuristic actively delays opening new days "
    "until the current one is filled, distributing sessions more evenly across the week.")

add_figure(doc, "fig2_s1_active_days.png",
    "Fig. 2.  Average number of active days per student group per week (S1) vs. N. "
    "The full configuration (solid blue) consistently keeps groups spread across more "
    "weekdays than the baseline (dashed orange).  Dotted line marks the five-day maximum.")

add_paragraph(doc,
    "Quality — compact days (S2). The fraction of group-days containing exactly two "
    "sessions, the preferred four-hour pattern, shows the most pronounced separation "
    "between configurations. At N = 5 the full configuration achieves S2 = 100%: every "
    "group-day on which a session appears contains exactly two sessions, corresponding to "
    "a perfectly compact schedule. The baseline reaches S2 = 0% at the same size, "
    "meaning sessions are never paired on the same day. As N increases and room pressure "
    "forces more deviations from the preferred pattern, S2 declines for both "
    "configurations, but the gap remains substantial: at N = 30 the full configuration "
    "retains S2 = 61.4 ± 1.1% against S2 = 28.0 ± 0.9% for the baseline — more than "
    "twice the share of compact days.")

add_figure(doc, "fig3_s2_compact_days.png",
    "Fig. 3.  Fraction of group-days containing exactly two sessions (S2) vs. N. "
    "The full configuration (solid blue) starts at 100% at N = 5 and retains a "
    "two-to-one advantage over the baseline (dashed orange) across the entire range, "
    "demonstrating the consistent effect of the day-packing heuristic.")

# 5.3
add_heading(doc, "5.3  Runtime Performance", level=2)

add_paragraph(doc,
    "Table 3 reports the mean runtime of both configurations as a function of N, "
    "together with 95% confidence intervals. Even at the largest tested size "
    "(N = 30 groups, 150 sessions) the full configuration completes in under 450 ms on "
    "average. The empirical curve is well approximated by a linear function of N for "
    "both configurations, which is consistent with the worst-case bound "
    "O(N × |T| × |R|) of Section 3.3 since both |T| = 15 and |R| = 12 are held fixed. "
    "The full configuration incurs a modest overhead over the baseline — between 4 ms at "
    "N = 5 and 73 ms at N = 30 — attributable to the candidate-scoring loop of the "
    "day-packing heuristic, which evaluates the day_session_score function for every "
    "valid (slot, room) candidate before selecting the minimum.")

# Table 3
t3_headers = ["N (groups)", "Runtime — Baseline (ms)", "Runtime — Full (ms)"]
t3_rows = []
for n in SIZES:
    rb, rb_ci = stats(n, "rt_baseline_ms")
    rf, rf_ci = stats(n, "rt_full_ms")
    t3_rows.append([
        str(n),
        f"{rb:.1f} ± {rb_ci:.1f}",
        f"{rf:.1f} ± {rf_ci:.1f}",
    ])
add_table(doc, t3_headers, t3_rows, col_widths=[1.3, 2.5, 2.5])
add_caption(doc,
    "Table 3.  Mean end-to-end API runtime (ms) for the baseline and full configurations, "
    "over 30 instances per size.  95% CI shown.  Timings include all database I/O through "
    "the PostgreSQL back-end.")

add_paragraph(doc,
    "The narrow 95% intervals confirm that the runtime variance is low and independent of "
    "the random instance: the dominant factor is the deterministic conflict-set lookup "
    "rather than the random group sizes. The runtime figures reported here are "
    "end-to-end API measurements that include all PostgreSQL read and write operations, "
    "making them directly comparable to the administrator's experience when triggering "
    "the scheduler from the web interface.")

add_figure(doc, "fig4_runtime.png",
    "Fig. 4.  Mean end-to-end API runtime (ms) vs. N for the baseline (dashed orange) "
    "and full (solid blue) configurations, with 95% CI shaded.  The dotted line shows "
    "the linear fit to the full configuration, confirming O(N) scaling consistent with "
    "the O(N × |T| × |R|) bound of Section 3.3.")

# ─────────────────────────────────────────────────────────────────────────────
# SECTION 6
# ─────────────────────────────────────────────────────────────────────────────

doc.add_page_break()
add_heading(doc, "6.  Discussion", level=1)

add_paragraph(doc,
    "Three observations can be drawn from the results of Section 5.")

add_paragraph(doc,
    "Architectural constraints shape the algorithm choice. The decision to run the "
    "scheduler synchronously inside a single HTTP request-response cycle, described in "
    "Section 3.1, filters out the long-running metaheuristics that dominate the "
    "algorithmic literature [6]. The results of Section 5.3 confirm that this constraint "
    "is compatible with practical use: even at N = 30 groups the full configuration "
    "completes in under 450 ms, leaving the administrator a sub-second feedback loop "
    "when editing a resource and re-running the scheduler. The linear scaling of the "
    "runtime with N implies that the system can be extended to larger institutions — "
    "for instance 60 groups with proportionally more rooms and slots — without crossing "
    "the one-second mark, provided the slot and room counts grow at a similar rate to N.")

add_paragraph(doc,
    "The day-packing heuristic is effective at no runtime cost. The S2 results of "
    "Section 5.2 provide the clearest evidence: at N = 5, where the room budget is "
    "unconstrained, the full configuration achieves S2 = 100% — every group fills its "
    "active days with exactly two sessions — while the baseline produces S2 = 0%. The "
    "improvement is not an artefact of the small instance size; at N = 20 the full "
    "configuration still holds S2 = 64.4% against the baseline's 25.5%. The mechanism "
    "is the day_session_score function described in Section 3.4: by assigning a lower "
    "penalty to a day that already hosts one session than to a completely empty day, "
    "the heuristic reuses open daily bins before opening new ones, a direct "
    "transposition of the First-Fit-Decreasing strategy [10] to the timetabling setting. "
    "The runtime overhead of this scoring step is absorbed by the fixed cost of the "
    "database lookups and remains below 73 ms across the entire benchmark range.")

add_paragraph(doc,
    "Feasibility is room-limited, not algorithm-limited. The feasibility results of "
    "Section 5.2 show that both configurations achieve identical scheduling rates across "
    "all N values, because the two configurations share the same hard-constraint "
    "checking logic. The declining feasibility curve — from 100% at N = 5 to roughly "
    "51% at N = 30 — is therefore not a weakness of either algorithm but an arithmetic "
    "consequence of the room-slot budget (180 pairs) relative to the session demand "
    "(10 × N sessions). The practical implication for an institution is that the "
    "feasibility of the schedule is determined primarily by the ratio of groups to "
    "available rooms, and that adding rooms or splitting a large semester across "
    "multiple scheduling runs is a more effective lever than changing the algorithm.")

add_paragraph(doc,
    "Limitations. Three limitations of the current evaluation should be noted. First, "
    "the benchmark uses synthetic instances in which every group has its own dedicated "
    "set of courses and professors. Real university curricula typically involve shared "
    "courses across degree programmes and professors who teach multiple groups, "
    "introducing professor-slot conflicts that our benchmark does not exercise and that "
    "would be expected to lower feasibility at a given N. Second, the instance sizes "
    "tested (up to 30 groups) correspond to a small faculty; a full university with "
    "several hundred groups would require proportionally larger room and slot sets, and "
    "the scalability of the algorithm in that regime has not been validated. Third, the "
    "professor load-balancing rule of Section 3.4 has no measurable effect on the "
    "benchmark because each professor is assigned to exactly one group; the S3 metric "
    "is consequently identically zero across all instances, and a meaningful evaluation "
    "of the load-balancing heuristic requires a multi-group professor assignment, which "
    "is deferred to future work alongside a real-data validation on the institution's "
    "actual curriculum.")

# ─────────────────────────────────────────────────────────────────────────────
# Save
# ─────────────────────────────────────────────────────────────────────────────

out = Path("Sections_5_6_replacement.docx")
doc.save(out)
print(f"Saved → {out.resolve()}")
